#!/usr/bin/env python3
import os
import re
from docx import Document
for folder, subfolders, files in os.walk(os.getcwd()):
    for file in files:
        if ".docx" in file:
            def docx_replace_regex(doc_obj, regex, replace):

                for p in doc_obj.paragraphs:
                    if regex.search(p.text):
                        inline = p.runs
                        # Loop added to work with runs (strings with same style)
                        for i in range(len(inline)):
                            if regex.search(inline[i].text):
                                text = regex.sub(replace, inline[i].text)
                                inline[i].text = text

                for table in doc_obj.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            docx_replace_regex(cell, regex, replace)


            regex1 = re.compile(r"Muk")
            replace1 = r" Fuk"
            filename = file
            doc = Document(filename)
            docx_replace_regex(doc, regex1, replace1)
            doc.save(filename)

        elif ".txt" in file:
            with open(file) as f:
                newText = f.read().replace('Muk', 'Fuk')
                with open(file, "w") as f:
                    f.write(newText)




