import time
import docx
import os
import xlrd
import keyboard
import webbrowser
import datetime
from datetime import date


# Current Directory
dir = os.getcwd()

# Excel Stuff
loc = dir + '/list.xlsx'
wb = xlrd.open_workbook(loc)
sheet = wb.sheet_by_index(0)

# Format Date For Letter
date = (time.strftime("%m/%d/%Y"))


def letter():
    # Make folders for Letters and Labels
    if not os.path.exists(dir + '/Letters'):
        os.mkdir(dir + '/Letters')
    body1 = "\t\t\t\t\t\t\t\t\t\t\t\t\t\t{}\nDear {} Management,\n\nThe Mukilteo Water and Wastewater District (District) is strengthening its responsibility of monitoring \
    wastewater pre-treatment requirements to businesses within its service area.  Pursuant to the District’s adopted \
    wastewater pre-treatment code:\n\n"
    bolditalic = "\tCode excerpts from Section I. page 25 and 26:\n\n\t3. All grease interceptors" \
                 " and/or mechanical devices must be pumped out \tcompletely once every three months, or more frequently, as required" \
                 " by the \tDistrict.\n\n\t4. Oil/water separators must be pumped out completely when the oil level on \ttop reaches 2 " \
                 "inches in thickness or the debris or sludge level occupies 25 \tpercent of the \tvolume, or more frequently, as " \
                 "required by the District.\n\n\t8. A log indicating each pumping of a grease interceptor for" \
                 " the previous 12 \tmonths shall be maintained by each food service establishment. A log indicating \teach pumping of" \
                 " an oil/water separator for the previous 12 months shall be \tmaintained by the Owner or his representative. This " \
                 "log shall include date, time, \tamount pumped, hauler and disposal site, and shall be kept in a conspicuous" \
                 "\tlocation for inspection by Health Department or District personnel. The \tmaintenance record log shall be recorded" \
                 " in the format on file with the District.\n\n\t9. Maintenance Reporting. The information required in the maintenance " \
                 "log \tshall be submitted to the District annually. The reporting period is January 1st \tthrough December 31st of " \
                 "each year. The report shall be submitted within 30 \tdays after the end of the reporting period.\n"
    body = "\nOur records indicate that we have not received maintenance records for the {} at {}\
     for the period January 1st through December 31st 2019. Please submit the maintenance records by\
     March 31, 2020. If the {} was not pumped in 2019, you will need to immediately schedule a pump-out and\
      provide the District with the pump-out schedule by March 31, 2020.  All correspondence, including the 2019 maintenance\
     records or 2020 pump-out schedule should be sent to the District at 7824 Mukilteo Speedway Mukilteo WA 98275, \
    attention Brandon Henson, or emailed to Brandonh@mukilteowwd.org. As a reminder, the log shall include Date, Time,\
    Amount Pumped, hauler, and Disposal Site. Lastly, the District will also begin conducting on-site inspections in the\
     very near future. Businesses that do not have current maintenance records for will be our first priority. Detailed\
     information regarding the District’s pre-treatment code can be found at the following web-site:\n\n\tComplete\
     pre-treatment code:\n\thttps://www.mukilteowwd.org/DocumentCenter/View/94/Pretreatment-PDF\n\n\tMaintenance log:\n\t\
    https://www.mukilteowwd.org/DocumentCenter/View/349/MWWD_Grease_Trap_Maintenance_Log\n\nShould you have any questions,\
    please do not hesitate to contact me at 425-355-3355 or email at brandonh@mukilteowwd.org.\n\nThank you for your time,\
           \n\nBrandon Henson\nMukilteo Water & Wastewater District"
    ind = 0
    for i in range(sheet.nrows):
        doc = docx.Document()
        name = sheet.cell_value(i, 0)
        address = sheet.cell_value(i, 1)
        typeof = sheet.cell_value(i, 2)
        ind = int(ind) + 1
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture('C:/Users/brandonh/Desktop/PretreatmentMaintenanceRecordsCode/1.JPG')
        p.add_run(body1.format(date, name))
        runner = p.add_run(bolditalic)
        p.add_run(body.format(typeof, address, typeof))
        runner.bold = True
        runner.italic = True
        doc.save(dir + '/Letters/' + str(name) + '.docx')


def envelopes():
    body = '\n\n\n\t\t\t\t\t\t{} Management\n\t\t\t\t\t\t{}\n\t\t\t\t\t\tMukilteo WA 98275'
    ind = 0
    if not os.path.exists(dir + '/labels'):
        os.mkdir(dir + '/Labels')
    for i in range(sheet.nrows):
        doc = docx.Document()
        name = sheet.cell_value(i, 0)
        address = sheet.cell_value(i, 1)
        ind = int(ind) + 1
        p = doc.add_paragraph()
        r = p.add_run()
        p.add_run(body.format(name, address))
        doc.save(dir + '/Labels/LABEL_' + str(name) + '.docx')


def consolidate():
    input('Press ENTER to consolidate')
    webbrowser.open('C:/ProgramData/Microsoft/Windows/Start Menu/Programs/Word 2016')
    time.sleep(3)
    keyboard.press_and_release('enter')
    time.sleep(1)
    keyboard.press_and_release('alt')
    time.sleep(1)
    keyboard.press_and_release('n')
    time.sleep(1)
    keyboard.press_and_release('j')
    time.sleep(1)
    keyboard.press_and_release('f')
    time.sleep(30)
    keyboard.press_and_release('alt')
    time.sleep(1)
    keyboard.press_and_release('p')
    time.sleep(1)
    keyboard.press_and_release('o')
    time.sleep(1)
    keyboard.press_and_release('down')
    time.sleep(1)
    keyboard.press_and_release('enter')
    time.sleep(1)
    keyboard.press_and_release('alt')
    time.sleep(1)
    keyboard.press_and_release('f')
    time.sleep(1)
    keyboard.press_and_release('p')
    time.sleep(1)
    keyboard.press_and_release('l')
    time.sleep(1)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('down')
    time.sleep(.25)
    keyboard.press_and_release('enter')
    time.sleep(1)
    keyboard.press_and_release('esc')
    time.sleep(1)
    keyboard.press_and_release('ctrl + a')
    time.sleep(1)
    keyboard.press_and_release('alt')
    time.sleep(1)
    keyboard.press_and_release('h')
    time.sleep(1)
    keyboard.press_and_release('f')
    time.sleep(1)
    keyboard.press_and_release('s')
    time.sleep(1)
    keyboard.press_and_release('1')
    time.sleep(1)
    keyboard.press_and_release('4')
    time.sleep(1)
    keyboard.press_and_release('enter')
    time.sleep(1)
    keyboard.press_and_release('alt')
    time.sleep(1)
    keyboard.press_and_release('f')
    time.sleep(1)
    keyboard.press_and_release('a')
    time.sleep(1)
    keyboard.press_and_release('o')
    time.sleep(3)
    keyboard.write(dir + '\\Labels\\COMBINED')
    time.sleep(1)
    keyboard.press_and_release('enter')
    time.sleep(1)


def turnedin():
    year = int(input("4 digit year"))
    month = int(input("1 or 2 digit month"))
    day = int(input("1 or 2 digit day"))
    datetostartfrom = datetime.datetime(year, month, day, 0, 0).timestamp()
    filelist = open(str(os.getcwd()) + '/SubmittedPretreatmentMaintenanceRecords.csv', 'a')
    count = 0
    for i in os.listdir('S:/Pretreatment/Pretreatment Inspections'):
        foldername = 'S:/Pretreatment/Pretreatment Inspections/' + i + '/' + 'MaintenanceRecords'
        for file in os.listdir(foldername):
            file1 = os.path.join("S:/Pretreatment/Pretreatment Inspections/" + i + '/' + 'MaintenanceRecords/' + file)
            ModTime = os.path.getmtime(file1)
            if int(ModTime) > int(datetostartfrom):
                filelist.writelines(str(file1))
                filelist.writelines('\n')
                count = count + 1
    filelist.writelines(str(count) + ' files')


def needstoturnin():
    ownerlist = ["H.P. AUTO BATH",
                 "KARMICHAEL AUTO SALON",
                 "LES SCHWAB",
                 "THE SPOT (CARWASH",
                 "MUKILTEO SHELL",
                 "SMART AUTO SERVICE",
                 "A BURGER PLACE",
                 "ARNIES RESTAURANT ",
                 "CAFÉ SOLEIL",
                 "AMICI",
                 "HILTON GARDEN INN",
                 "IVARS FISH BAR",
                 "KAMIAK HIGH SCHOOL",
                 "KISS OF THE OCEAN",
                 "LOMBARDOS",
                 "MONGOLIAN GRILL",
                 "MUKILTEO SPEEDWAY CAFE",
                 "SAKUMA",
                 "SUBWAY",
                 "TACO BELL",
                 "WHIDBEY COFFEE ",
                 "AMBROSIA",
                 "AZTECA",
                 "BITE OF NEW YORK",
                 "BROOKLYN BROTHERS PIZZA",
                 "DIAMOND KNOT BREWERY",
                 "DIAMOND KNOT BREWING CO",
                 "GARLIC JIMS FAMOUS GOURMET PIZZA",
                 "GROUCHY CHEF",
                 "GYRO STOP",
                 "H.P. MIDDLE SCHOOL",
                 "H.P. RETIREMENT",
                 "HANAMI",
                 "HENRYS DONUTS",
                 "JOHNS GRILL",
                 "KAMIAK HIGH SCHOOL",
                 "KINGS TERIYAKI",
                 "KOSTAS",
                 "LOTUS",
                 "MUKILTEO CHOCOLATE CO",
                 "MUKILTEO LODGE",
                 "MUKILTEO THAI",
                 "PATTY'S EGG NEST",
                 "PIZZA HUT",
                 "RAINBOW WOK & TERIYAKI ",
                 "ROSEHILL COMMUNITY CENTER",
                 "SABOR A MEXICO ",
                 "SHAKE 'N GO",
                 "STARBUCKS COFFEE",
                 "STAYBRIDGE SUITES",
                 "SULLYS",
                 "THE SCOTSMAN",
                 "THE SYDNEY BAKERY",
                 "TRAXX ",
                 "WELLERS",
                 "ZS BURGER ",
                 "MUKILTEO SHELL",
                 "KARMICHAEL AUTO SALON ",
                 "SUGARCANE"]
    year = int(input("4 digit year"))
    month = int(input("1 or 2 digit month"))
    day = int(input("1 or 2 digit day"))
    datetostartfrom = datetime.datetime(year, month, day, 0, 0).timestamp()
    results = open('NeedsMaintenanceRecords.csv', 'a')
    list = []
    count = 0
    for i in os.listdir('S:/Pretreatment/Pretreatment Inspections'):
        x = i.split('_')
        name = x[0]
        # address = x[1]
        if name.upper() in ownerlist:
            list.append(name)
            foldername = 'S:/Pretreatment/Pretreatment Inspections/' + i + '/' + 'MaintenanceRecords'
            if not os.path.exists(foldername):
                os.mkdir(foldername)
            for file in os.listdir(foldername):
                file1 = os.path.join(
                    "S:/Pretreatment/Pretreatment Inspections/" + i + '/' + 'MaintenanceRecords/' + file)
                ModTime = os.path.getmtime(file1)
                if int(ModTime) > int(datetostartfrom):
                    list.remove(name)
                    # print(name)
    for i in list:
        results.write(str(i) + '\n')


def menu_selection(prompt, dispatch_dict):
    while True:
        response = input(prompt)
        try:
            if dispatch_dict[response]() == "Exit Menu":
                break
        except KeyError:
            print("\nPick from the listed options.")


prompt = ('\nSelect an option:\n'
          '[1] Needs To Turn In Maintenance Records\n'
          '[2] Already Turned In Maintenance Records\n'
          '[3] Make Letters\n'
          '[4] Make Envelopes\n'
          '[5] Consolidate Envelopes Doc Into 1\n'
          '[6] Exit\n')


def exit():
    return "Exit Menu"


arg_dict = {"1": needstoturnin, "2": turnedin, "3": letter, "4": envelopes, "5": consolidate, "6": exit}

if __name__ == '__main__':
    menu_selection(prompt, arg_dict)
