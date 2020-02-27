import time
import docx

body = "{}\nDear {} Management,\nOur records indicate that we have not received maintenance records for the {} at {}." \
       " Our pretreatment code requires customers with these devices to maintain them at a minimum of every three months" \
       " and report all maintenance to the District. We will be conducting on-site inspections in the very near future." \
       "  Businesses that we do not have current maintenance records for will be our first priority. You can email your " \
       "records to: Brandonh@mukilteowwd.org or return them to 7824 Mukilteo Speedway Mukilteo WA 98275 C\O " \
       "Brandon Henson. The log shall include Date, Time, Amount Pumped, hauler, and Disposal Site.\n\n"

bolditalic = "Code excerpts from Section I. page 25 and 26:\n3. All grease interceptors" \
             " and/or mechanical device must be pumped out completely once every three months, or more frequently, as required" \
             " by the District.\n\n4. Oil/water separators must be pumped out completely when the oil level on top reaches 2 " \
             "inches in thickness or the debris or sludge level occupies 25 percent of the volume, or more frequently, as " \
             "required by the District. \n\n8. A log indicating each pumping of a grease interceptor for" \
             " the previous 12 months shall be maintained by each food service establishment. A log indicating each pumping of" \
             " an oil/water separator for the previous 12 months shall be maintained by the Owner or his representative. This " \
             "log shall include date, time, amount pumped, hauler and disposal site, and shall be kept in a conspicuous " \
             "location for inspection by Health Department or District personnel. The maintenance record log shall be recorded" \
             " in the format on file with the District.\n\n9. Maintenance Reporting. The information required in the maintenance " \
             "log shall be submitted to the District annually. The reporting period is January 1st through December 31st of " \
             "each year. The report shall be submitted within 30 days after the end of the reporting period."

footer = "\n\nFor the complete code, please visit: https://www.mukilteowwd.org/DocumentCenter/View/94/Pretreatment-PDF\n\n" \
         "For a maintenance log, please visit: https://www.mukilteowwd.org/DocumentCenter/View/349/MWWD_Grease_Trap_Maintenance_Log\n\n" \
         "Thank you for your time,\nBrandon Henson" \
         "\nMukilteo Water & Wastewater District\n425-355-3355"

# mm/yyyy format
date = (time.strftime("%m/%d/%Y"))
customer = [
    ["Ambrosia", "415 Lincoln Ave", "Grease Trap"],
    ["Azteca", "11811 Mukilteo Speedway", "Grease Trap"],
    ["Bite of New York", "10100 Mukilteo Speedway", "Grease Trap"],
    ["Brooklyn Bros. Pizza", "8326 Mukilteo Speedway", "Grease Trap"],
    ["Charles At Smugglers", "8340 53rd Ave W", "Grease Trap"],
    ["Diamond Knot", "621 Front St", "Grease Trap"],
    ["Garlic Jim's", "10924 Mukilteo Speedway", "Grease Trap"],
    ["Grouchy Chef", "4433 Russell Rd", "Grease Trap"],
    ["Gyro Stop", "11811 Mukilteo Speedway", "Grease Trap"],
    ["H.P. Middle School", "5000 Harbour Pointe Blvd", "Grease Trap"],
    ["H.P. Retirement", "10200 Harbour Place", "Grease Trap"],
    ["Hanami", "11811 Mukilteo Speedway", "Grease Trap"],
    ["Henry's Donuts", "10924 Mukilteo Speedway", "Grease Trap"],
    ["John's Grill", "649 5th St", "Grease Trap"],
    ["Kamiak High School", "10801 Harbour Pt Blvd N", "Grease Trap"],
    ["King's Teriyaki", "8229 44th Ave W", "Grease Trap"],
    ["Kosta's", "8306 Mukilteo Speedway", "Grease Trap"],
    ["Lotus", "10924 Mukilteo Speedway", "Grease Trap"],
    ["Mukilteo Chocolate Co.", "407 Lincoln Ave", "Grease Trap"],
    ["Mukilteo Lodge", "7928 Mukilteo Speedway", "Grease Trap"],
    ["Mukilteo Thai", "8410 Mukilteo Speedway", "Grease Trap"],
    ["Patty's Egg Nest", "8204 Mukilteo Speedway", "Grease Trap"],
    ["Pizza Hut", "11707 Mukilteo Speedway", "Grease Trap"],
    ["Rainbow Wok & Teriyaki", "8410 Mukilteo Speedway", "Grease Trap"],
    ["Rosehill Community Center", "304 Lincoln Ave", "Grease Trap"],
    ["Sabor A Mexico", "8410 Mukilteo Speedway", "Grease Trap"],
    ["Shake 'n Go", "9999 Harbour Place", "Grease Trap"],
    ["Starbucks", "10100 Mukilteo Speedway", "Grease Trap"],
    ["Staybridge Suites", "9600 Harbour Place", "Grease Trap"],
    ["Sully's", "403 Lincoln Ave", "Grease Trap"],
    ["The Scotsman", "11601 Harbour Pointe Blvd", "Grease Trap"],
    ["The Sydney Bakery", "613 5th St", "Grease Trap"],
    ["Traxx", "4329 Chennault Beach Rd", "Grease Trap"],
    ["Weller's", "8490 Mukilteo Speedway", "Grease Trap"],
    ["Z's Burgers", "8307 Mukilteo Speedway", "Grease Trap"],
    ["A Burger Place", "11601 Harbour Pt Blvd", "Grease Interceptor"],
    ["Arnie's", "714 2nd St", "Grease Interceptor"],
    ["Café Soleil", "9999 Harbour Place", "Grease Interceptor"],
    ["Fra Amici", "8004 Mukilteo Speedway", "Grease Interceptor"],
    ["Hilton Garden Inn", "8401 Paine Field Blvd", "Grease Interceptor"],
    ["Ivar's", "710 Front St", "Grease Interceptor"],
    ["Kamiak High School", "10801 Harbour Pt Blvd N", "Grease Interceptor"],
    ["Kiss Of The Ocean", "801 2nd St", "Grease Interceptor"],
    ["Lombardo's Pizza", "10809 Mukilteo Speedway", "Grease Interceptor"],
    ["Mongolian Grill", "11700 Mukilteo Speedway", "Grease Interceptor"],
    ["Mukilteo Shell", "10900 Mukilteo Speedway", "Grease Interceptor"],
    ["Mukilteo's Spdwy Cafe", "11707 Mukilteo Speedway", "Grease Interceptor"],
    ["Sakuma", "10924 Mukilteo Speedway", "Grease Interceptor"],
    ["Subway", "9999 Harbour Place", "Grease Interceptor"],
    ["Taco Bell", "8401 Mukilteo Speedway", "Grease Interceptor"],
    ["Whidbey Island Coffee", "204 Lincoln Ave", "Grease Interceptor"],
    ["H.P. Auto Bath", "11400 Mukilteo Speedway", "Oil \ Water Seperator"],
    ["Karmichael Auto Salon", "11400 'B' Mukilteo Speedway", "Oil \ Water Seperator"],
    ["Les Schwab", "11500 Mukilteo Speedway", "Oil \ Water Seperator"],
    ["The Spot", "8203 Mukilteo Speedway", "Oil \ Water Seperator"],
    ["Mukilteo Shell", "10900 Mukilteo Speedway", "Oil \ Water Seperator"],
    ["Smart Auto Service", "11338 Mukilteo Speedway", "Oil \ Water Seperator"],
    ["Sugar Cane Ice cream", "8410 Mukilteo Speedway G", "Grease Trap"]
]
ind = 0
for i in customer:
    doc = docx.Document()
    name = (customer[ind][0])
    address = (customer[ind][1])
    typeof = (customer[ind][2])
    ind = int(ind) + 1
    p = doc.add_paragraph()
    p.add_run(body.format(date, name, typeof, address))
    runner = p.add_run(bolditalic)
    p.add_run(footer)
    runner.bold = True
    runner.italic = True
    doc.save(str(name) + '.docx')
